from flask import Flask, render_template, request, jsonify, send_file, redirect, url_for, session
from flask_cors import CORS
from flask_sqlalchemy import SQLAlchemy
from flask_login import LoginManager, UserMixin, login_user, login_required, logout_user, current_user
from werkzeug.security import generate_password_hash, check_password_hash
import os
import shutil
import pytesseract
from PIL import Image
import PyPDF2
import docx
# import speech_recognition as sr
# from pydub import AudioSegment
import logging
import io
import requests
from bs4 import BeautifulSoup
from youtube_transcript_api import YouTubeTranscriptApi
import yt_dlp
import google.generativeai as genai
import time
import networkx as nx
import matplotlib
matplotlib.use('Agg')  # Use non-GUI backend
import matplotlib.pyplot as plt
from gtts import gTTS
from werkzeug.utils import secure_filename
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.units import inch
import base64
from dotenv import load_dotenv
import glob
from datetime import datetime, timedelta
from groq import Groq
import secrets
import hashlib
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from functools import wraps
from collections import defaultdict

load_dotenv()
if os.getenv('GEMINI_API_KEY'):
    genai.configure(api_key=os.getenv('GEMINI_API_KEY'))

app = Flask(__name__)
CORS(app)
app.config['SECRET_KEY'] = os.getenv('SECRET_KEY', 'your-secret-key-change-this-in-production')
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:////tmp/users.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['UPLOAD_FOLDER'] = '/tmp/Uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024
app.config['SESSION_COOKIE_SAMESITE'] = 'Lax'
app.config['SESSION_COOKIE_SECURE'] = False  # Set to True in production with HTTPS
app.config['PERMANENT_SESSION_LIFETIME'] = timedelta(days=7)
app.config['SESSION_COOKIE_HTTPONLY'] = True

# Initialize extensions
db = SQLAlchemy(app)
login_manager = LoginManager(app)
login_manager.login_view = 'login_page'

@login_manager.unauthorized_handler
def unauthorized():
    # Return JSON for API/AJAX requests, redirect for page requests
    if request.is_json or request.headers.get('Content-Type') == 'application/json' or \
       request.headers.get('X-Requested-With') == 'XMLHttpRequest':
        return jsonify({'error': 'Authentication required', 'redirect': '/login'}), 401
    return redirect(url_for('login_page'))

# User Model
class User(UserMixin, db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(100), nullable=False)
    email = db.Column(db.String(120), unique=True, nullable=False)
    password = db.Column(db.String(200), nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

# OTP Model for password reset
class PasswordResetOTP(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    email = db.Column(db.String(120), nullable=False)
    otp_hash = db.Column(db.String(200), nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    expires_at = db.Column(db.DateTime, nullable=False)
    is_used = db.Column(db.Boolean, default=False)
    
# Rate limiting storage
otp_request_tracker = defaultdict(list)  # email -> [timestamps]

@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))

# Create database tables
with app.app_context():
    db.create_all()

if not os.path.exists(app.config['UPLOAD_FOLDER']):
    os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
if not os.path.exists('static'):
    os.makedirs('static', exist_ok=True)

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)
# Cleanup old files in static folder
def cleanup_old_files():
    try:
        cutoff_time = datetime.now() - timedelta(hours=1)
        static_dir = '/tmp' if os.environ.get('VERCEL') else 'static'
        patterns = [
            f'{static_dir}/audio_*.mp3', f'{static_dir}/video_*.png',
            f'{static_dir}/mindmap_*.png', f'{static_dir}/report_*.pdf',
            f'{static_dir}/summary_*.*'
        ]
        for pattern in patterns:
            for filepath in glob.glob(pattern):
                try:
                    file_time = datetime.fromtimestamp(os.path.getmtime(filepath))
                    if file_time < cutoff_time:
                        os.remove(filepath)
                        logger.info(f"Cleaned up old file: {filepath}")
                except Exception as e:
                    logger.error(f"Error deleting {filepath}: {e}")
    except Exception as e:
        logger.error(f"Cleanup error: {e}")

# Run cleanup on startup
cleanup_old_files()

SUPPORTED_LANGUAGES = {'en', 'es', 'fr', 'de', 'zh-cn', 'ja', 'ru', 'ar', 'hi', 'pt'}

pytesseract.pytesseract.tesseract_cmd = shutil.which('tesseract') or r'C:\Program Files\Tesseract-OCR\tesseract.exe'

def extract_text_from_file(file_path, file_type, language='en'):
    try:
        if file_type == 'pdf':
            try:
                with open(file_path, 'rb') as f:
                    pdf = PyPDF2.PdfReader(f)
                    if len(pdf.pages) == 0:
                        logger.error(f"PDF file {file_path} is empty")
                        return ''
                    text = ''
                    for page in pdf.pages:
                        extracted = page.extract_text() or ''
                        text += extracted
                    if not text.strip():
                        logger.error(f"No text extracted from PDF {file_path}")
                        return ''
                    return text
            except PyPDF2.errors.PdfReadError as e:
                logger.error(f"PDF read error for {file_path}: {e}")
                return ''
            except Exception as e:
                logger.error(f"Unexpected error processing PDF {file_path}: {e}")
                return ''
        elif file_type in ['txt', 'text', 'md', 'markdown']:
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    text = f.read()
                if not text.strip():
                    logger.error(f"No text in file {file_path}")
                    return ''
                return text
            except Exception as e:
                logger.error(f"Error reading text file {file_path}: {e}")
                return ''
        elif file_type == 'docx':
            try:
                doc = docx.Document(file_path)
                text = '\n'.join([para.text for para in doc.paragraphs if para.text.strip()])
                if not text:
                    logger.error(f"No text extracted from DOCX {file_path}")
                    return ''
                return text
            except Exception as e:
                logger.error(f"Error processing DOCX {file_path}: {e}")
                return ''
        elif file_type in ['jpg', 'jpeg', 'png']:
            try:
                image = Image.open(file_path)
                text = pytesseract.image_to_string(image, lang=language)
                if not text.strip():
                    logger.error(f"No text extracted from image {file_path}")
                    return ''
                return text
            except Exception as e:
                logger.error(f"Error processing image {file_path}: {e}")
                return ''
        elif file_type in ['mp4', 'mp3']:
            logger.error(f"Audio/video processing not available in Python 3.13+")
            return ''
            # Audio processing disabled due to Python 3.13 compatibility
        else:
            logger.error(f"Unsupported file type: {file_type}")
            return ''
    except Exception as e:
        logger.error(f"General error extracting text from {file_path}: {e}")
        return ''

def extract_youtube_transcript(url, language='en'):
    try:
        video_id = url.split('v=')[1].split('&')[0]
        transcript = YouTubeTranscriptApi.get_transcript(video_id, languages=[language])
        text = ' '.join([entry['text'] for entry in transcript])
        if not text.strip():
            logger.error(f"No transcript extracted from YouTube URL {url}")
            return ''
        return text
    except Exception as e:
        logger.warning(f"Primary YouTube transcript extraction failed for {url}: {e}")
        try:
            ydl_opts = {'format': 'bestaudio/best', 'noplaylist': True}
            with yt_dlp.YoutubeDL(ydl_opts) as ydl:
                info = ydl.extract_info(url, download=False)
                temp_path = os.path.join(app.config['UPLOAD_FOLDER'], 'temp_audio.' + info['ext'])
                ydl.download([url])
                text = extract_text_from_file(temp_path, info['ext'], language)
                os.remove(temp_path) if os.path.exists(temp_path) else None
                if not text.strip():
                    logger.error(f"No text extracted from downloaded YouTube audio {url}")
                    return ''
                return text
        except Exception as e2:
            logger.error(f"YouTube fallback extraction error for {url}: {e2}")
            return ''

def generate_gemini_summary(text, language='en'):
    try:
        if not text.strip():
            logger.error("No text provided for summary")
            return ''
        
        # Map language codes to full language names for better AI understanding
        language_names = {
            'en': 'English',
            'es': 'Spanish',
            'fr': 'French',
            'de': 'German',
            'zh': 'Chinese',
            'zh-cn': 'Chinese',
            'ja': 'Japanese',
            'ru': 'Russian',
            'ar': 'Arabic',
            'hi': 'Hindi',
            'pt': 'Portuguese'
        }
        
        language_name = language_names.get(language.lower(), 'English')
        
        # Try Groq first (FASTEST!)
        groq_api_key = os.getenv('GROQ_API_KEY')
        if groq_api_key:
            try:
                client = Groq(api_key=groq_api_key)
                
                chat_completion = client.chat.completions.create(
                    messages=[
                        {
                            "role": "user",
                            "content": f"Summarize in {language_name}:\n{text[:1000]}",
                        }
                    ],
                    model="llama-3.1-8b-instant",
                    temperature=0.1,
                    max_tokens=200,  # Very short
                    stream=False,
                )
                
                summary = chat_completion.choices[0].message.content
                if summary.strip():
                    logger.info(f"Groq API generated summary in {language_name}")
                    return summary
            except Exception as e:
                logger.error(f"Groq API error: {e}")
        
        # Fallback to Gemini if available
        if os.getenv('GEMINI_API_KEY'):
            model = genai.GenerativeModel('gemini-1.5-flash')
            prompt = f"""Please provide a comprehensive summary of the following text. 
IMPORTANT: Write the entire summary in {language_name} language only.
Do not use English or any other language. Use only {language_name}.

Text to summarize:
{text}

Summary in {language_name}:"""
            
            response = model.generate_content(prompt)
            summary = response.text
            if not summary.strip():
                logger.error("Gemini returned empty summary")
                return ' '.join(text.split()[:100]) + '...' if len(text.split()) > 100 else text
            return summary
        
        # Final fallback - just return first 100 words
        logger.warning("No API key configured for summarization")
        words = text.split()[:100]
        return ' '.join(words) + '...' if len(words) > 100 else text
        
    except Exception as e:
        logger.error(f"Summary generation error: {e}")
        return ' '.join(text.split()[:100]) + '...' if len(text.split()) > 100 else text

def generate_gemini_answer(text, question, language='en', summary=None):
    try:
        if not text.strip() or not question.strip():
            logger.error("Text or question missing for answer")
            return f"Error: Text or question missing in {language}"
        
        # Map language codes to full language names
        language_names = {
            'en': 'English',
            'es': 'Spanish',
            'fr': 'French',
            'de': 'German',
            'zh': 'Chinese',
            'zh-cn': 'Chinese',
            'ja': 'Japanese',
            'ru': 'Russian',
            'ar': 'Arabic',
            'hi': 'Hindi',
            'pt': 'Portuguese'
        }
        
        language_name = language_names.get(language.lower(), 'English')
        
        # Try Groq first (faster and better)
        groq_api_key = os.getenv('GROQ_API_KEY')
        if groq_api_key:
            try:
                client = Groq(api_key=groq_api_key)
                # Use summary if available (much faster), otherwise use limited text
                context = summary if summary else text[:1000]
                
                chat_completion = client.chat.completions.create(
                    messages=[
                        {
                            "role": "user",
                            "content": f"{context[:500]}\n\nQ: {question}\nA:",
                        }
                    ],
                    model="llama-3.1-8b-instant",
                    temperature=0.1,
                    max_tokens=100,  # Very short
                    stream=False,
                )
                
                answer = chat_completion.choices[0].message.content
                if answer.strip():
                    logger.info(f"Groq API answered in {chat_completion.usage.total_tokens} tokens")
                    return answer
            except Exception as e:
                logger.error(f"Groq API error: {e}")
        
        # Fallback to Gemini
        if os.getenv('GEMINI_API_KEY'):
            model = genai.GenerativeModel('gemini-1.5-flash')
            context = summary if summary else text[:1000]
            prompt = f"""Based on this context: {context}

Question: {question}

IMPORTANT: Provide your answer in {language_name} language only. Do not use English or any other language.

Answer in {language_name}:"""
            response = model.generate_content(prompt)
            answer = response.text
            if answer.strip():
                return answer
        
        # Final fallback
        logger.error("No API key configured")
        return f"I need an API key to answer questions. Please configure GROQ_API_KEY or GEMINI_API_KEY in the .env file."
    except Exception as e:
        logger.error(f"Answer generation error: {e}")
        return f"Error generating answer: {str(e)}"

def generate_mind_map(text):
    try:
        if not text.strip():
            logger.error("No text provided for mind map")
            return None
        
        concepts = None

        # Use Groq first for faster concept extraction
        groq_api_key = os.getenv('GROQ_API_KEY')
        if groq_api_key:
            try:
                client = Groq(api_key=groq_api_key)
                chat_completion = client.chat.completions.create(
                    messages=[
                        {
                            "role": "user",
                            "content": f"Extract exactly 6 key concepts from this text as a comma-separated list. Only return the concepts, nothing else:\n{text[:600]}",
                        }
                    ],
                    model="llama-3.1-8b-instant",
                    temperature=0.1,
                    max_tokens=60,
                )
                raw = chat_completion.choices[0].message.content
                concepts = [c.strip() for c in raw.split(',') if c.strip()][:6]
            except Exception as e:
                logger.error(f"Groq error in mind map: {e}")
                concepts = None

        if not concepts and os.getenv('GEMINI_API_KEY'):
            try:
                model = genai.GenerativeModel('gemini-1.5-flash')
                prompt = f"Extract exactly 6-8 key concepts from this text as a comma-separated list. Only return the concepts, nothing else: {text[:1000]}"
                response = model.generate_content(prompt)
                concepts = [c.strip() for c in response.text.split(',') if c.strip()][:8]
            except Exception as e:
                logger.error(f"Gemini error in mind map: {e}")
                concepts = None

        # Robust fallback — extract meaningful phrases from the text
        if not concepts or len(concepts) < 2:
            # Split into sentences, pick first word(s) of each as concept
            import re
            sentences = re.split(r'[.!?;]', text)
            concepts = []
            for s in sentences:
                words = s.strip().split()
                if len(words) >= 2:
                    concepts.append(' '.join(words[:3]))
                if len(concepts) >= 6:
                    break
            # If still not enough, just use individual significant words
            if len(concepts) < 2:
                words = [w for w in text.split() if len(w) > 4][:12]
                concepts = list(dict.fromkeys(words))[:6]  # deduplicate

        # Ensure we have at least 2 concepts
        if len(concepts) < 2:
            concepts = ['Main Topic', 'Key Concept', 'Supporting Idea']

        # Truncate long concept strings
        concepts = [c[:30] for c in concepts[:7]]
        
        # Create hierarchical mind map
        G = nx.DiGraph()
        main_topic = concepts[0]
        G.add_node(main_topic)
        
        for i, concept in enumerate(concepts[1:], 1):
            G.add_node(concept)
            if i <= 3:
                G.add_edge(main_topic, concept)
            else:
                parent = concepts[((i-1) % 3) + 1]
                G.add_edge(parent, concept)
        
        plt.figure(figsize=(8, 6), facecolor='white', dpi=80)
        pos = nx.spring_layout(G, k=2.5, iterations=20, seed=42)
        
        node_colors = ['#4e7ae7' if node == main_topic else '#7a4ee7' for node in G.nodes()]
        nx.draw_networkx_nodes(G, pos, node_color=node_colors, node_size=2500, alpha=0.92)
        nx.draw_networkx_edges(G, pos, edge_color='#aaaacc', arrows=True, arrowsize=15, width=1.5)
        nx.draw_networkx_labels(G, pos, font_size=7, font_weight='bold', font_color='white')
        
        plt.axis('off')
        plt.tight_layout(pad=0.5)
        
        img_buffer = io.BytesIO()
        plt.savefig(img_buffer, format='png', dpi=80, bbox_inches='tight', facecolor='white')
        img_buffer.seek(0)
        plt.close()
        
        img_base64 = base64.b64encode(img_buffer.getvalue()).decode('utf-8')
        return img_base64
    except Exception as e:
        logger.error(f"Mind map generation error: {e}")
        return None

def generate_audio_overview(text, language='en'):
    try:
        if not text.strip():
            logger.error("No text provided for audio generation")
            return jsonify({'error': 'No text provided for audio'}), 400
        
        # Map language codes for gTTS compatibility
        lang_map = {
            'en': 'en',
            'es': 'es',
            'fr': 'fr',
            'de': 'de',
            'zh': 'zh-CN',
            'zh-cn': 'zh-CN',
            'ja': 'ja',
            'ru': 'ru',
            'ar': 'ar',
            'hi': 'hi',
            'pt': 'pt'
        }
        
        gtts_lang = lang_map.get(language.lower(), 'en')
        
        # Create audio in memory buffer (no static folder)
        logger.info(f"Generating audio with language: {gtts_lang}")
        tts = gTTS(text=text, lang=gtts_lang, slow=False)
        
        # Save to memory buffer
        audio_buffer = io.BytesIO()
        tts.write_to_fp(audio_buffer)
        audio_buffer.seek(0)
        
        # Convert to base64 for embedding
        audio_base64 = base64.b64encode(audio_buffer.getvalue()).decode('utf-8')
        
        logger.info(f"Audio generated successfully")
        return jsonify({
            'success': True, 
            'audio_base64': audio_base64,
            'filename': f'audio_overview_{int(time.time())}.mp3'
        })
    except Exception as e:
        logger.error(f"Audio generation error: {e}")
        return jsonify({'error': f'Audio generation failed: {str(e)}'}), 500

def generate_video_overview(text, language='en'):
    try:
        if not text.strip():
            logger.error("No text provided for video generation")
            return jsonify({'error': 'No text provided'}), 400
        
        # Map language codes to full language names
        language_names = {
            'en': 'English',
            'es': 'Spanish',
            'fr': 'French',
            'de': 'German',
            'zh': 'Chinese',
            'zh-cn': 'Chinese',
            'ja': 'Japanese',
            'ru': 'Russian',
            'ar': 'Arabic',
            'hi': 'Hindi',
            'pt': 'Portuguese'
        }
        
        language_name = language_names.get(language.lower(), 'English')
        
        # Generate key points for video using Groq first
        groq_api_key = os.getenv('GROQ_API_KEY')
        if groq_api_key:
            try:
                client = Groq(api_key=groq_api_key)
                chat_completion = client.chat.completions.create(
                    messages=[
                        {
                            "role": "user",
                            "content": f"3 points:\n{text[:600]}",
                        }
                    ],
                    model="llama-3.1-8b-instant",
                    temperature=0.1,
                    max_tokens=100,
                )
                response_text = chat_completion.choices[0].message.content
                key_points = [line.strip() for line in response_text.split('\n') if line.strip() and len(line.strip()) > 3][:3]
            except Exception as e:
                logger.error(f"Groq error in video generation: {e}")
                key_points = text.split('.')[:5]
        elif os.getenv('GEMINI_API_KEY'):
            model = genai.GenerativeModel('gemini-1.5-flash')
            prompt = f"Extract 5 key points from this text as a numbered list in {language_name}: {text}"
            response = model.generate_content(prompt)
            key_points = [line.strip() for line in response.text.split('\n') if line.strip()][:5]
        else:
            sentences = text.split('.')[:5]
            key_points = [s.strip() for s in sentences if s.strip()]
        
        fig, ax = plt.subplots(figsize=(6, 5), facecolor='#1a1a2e', dpi=60)
        ax.set_xlim(0, 10)
        ax.set_ylim(0, 10)
        ax.axis('off')
        
        ax.text(5, 8.5, 'Video Overview', fontsize=16, weight='bold', 
                color='white', ha='center', va='center')
        
        y_position = 7
        for i, point in enumerate(key_points[:3], 1):
            point_text = point[:60] + '...' if len(point) > 60 else point
            ax.text(5, y_position, f"{i}. {point_text}", fontsize=9, 
                   color='#e0e0e0', ha='center', va='center')
            y_position -= 1.5
        
        plt.tight_layout(pad=0)
        
        img_buffer = io.BytesIO()
        plt.savefig(img_buffer, format='png', dpi=60, bbox_inches='tight', facecolor='#1a1a2e')
        img_buffer.seek(0)
        plt.close()
        
        # Convert to base64 for embedding
        img_base64 = base64.b64encode(img_buffer.getvalue()).decode('utf-8')
        
        return jsonify({
            'success': True, 
            'video_base64': img_base64,
            'key_points': key_points,
            'filename': f'video_overview_{int(time.time())}.png'
        })
    except Exception as e:
        logger.error(f"Video generation error: {e}")
        return jsonify({'error': 'Video generation failed'}), 500

def generate_report_document(text, language='en'):
    try:
        if not text.strip():
            logger.error("No text provided for report")
            return jsonify({'error': 'No text provided'}), 400
        
        # Map language codes
        language_names = {
            'en': 'English', 'es': 'Spanish', 'fr': 'French', 'de': 'German',
            'zh': 'Chinese', 'zh-cn': 'Chinese', 'ja': 'Japanese', 'ru': 'Russian',
            'ar': 'Arabic', 'hi': 'Hindi', 'pt': 'Portuguese'
        }
        language_name = language_names.get(language.lower(), 'English')
        
        # Use Groq first for faster report generation
        groq_api_key = os.getenv('GROQ_API_KEY')
        if groq_api_key:
            try:
                import json
                client = Groq(api_key=groq_api_key)
                chat_completion = client.chat.completions.create(
                    messages=[
                        {
                            "role": "user",
                            "content": f"JSON report:\n{text[:600]}\n{{\"title\":\"...\",\"intro\":\"...\",\"points\":[\"...\"],\"conclusion\":\"...\"}}",
                        }
                    ],
                    model="llama-3.1-8b-instant",
                    temperature=0.1,
                    max_tokens=150,
                )
                report_data = json.loads(chat_completion.choices[0].message.content.strip())
            except Exception as e:
                logger.error(f"Groq error in report: {e}")
                report_data = {
                    'title': 'Summary Report',
                    'introduction': text[:200] + '...',
                    'key_points': text.split('.')[:5],
                    'conclusion': 'This report summarizes the key information.'
                }
        elif os.getenv('GEMINI_API_KEY'):
            model = genai.GenerativeModel('gemini-1.5-flash')
            prompt = f"""Create a structured report from this text in {language_name}. Format as JSON with these keys:
            - title: A concise title
            - introduction: 2-3 sentence introduction
            - key_points: Array of 5-7 main points
            - conclusion: 2-3 sentence conclusion
            
            Text: {text}"""
            
            response = model.generate_content(prompt)
            try:
                import json
                report_data = json.loads(response.text.replace('```json', '').replace('```', '').strip())
            except:
                report_data = {
                    'title': 'Summary Report',
                    'introduction': text[:200] + '...',
                    'key_points': text.split('.')[:5],
                    'conclusion': text[-200:] if len(text) > 200 else text
                }
        else:
            report_data = {
                'title': 'Summary Report',
                'introduction': text[:200] + '...',
                'key_points': text.split('.')[:5],
                'conclusion': 'This report summarizes the key information from the provided content.'
            }
        
        # Create HTML report
        html_report = f"""
        <div style="font-family: Arial, sans-serif; max-width: 800px; margin: 0 auto; padding: 20px;">
            <h1 style="color: #2c3e50; border-bottom: 3px solid #4e7ae7; padding-bottom: 10px;">
                {report_data.get('title', 'Summary Report')}
            </h1>
            
            <h2 style="color: #4e7ae7; margin-top: 30px;">Introduction</h2>
            <p style="line-height: 1.8; color: #333;">
                {report_data.get('introduction', '')}
            </p>
            
            <h2 style="color: #4e7ae7; margin-top: 30px;">Key Points</h2>
            <ul style="line-height: 2; color: #333;">
                {''.join([f'<li>{point}</li>' for point in report_data.get('key_points', [])])}
            </ul>
            
            <h2 style="color: #4e7ae7; margin-top: 30px;">Conclusion</h2>
            <p style="line-height: 1.8; color: #333;">
                {report_data.get('conclusion', '')}
            </p>
            
            <div style="margin-top: 40px; padding-top: 20px; border-top: 1px solid #ddd; color: #999; font-size: 12px;">
                Generated by Smart Summarizer
            </div>
        </div>
        """
        
        return jsonify({
            'success': True,
            'report_html': html_report,
            'report_data': report_data
        })
    except Exception as e:
        logger.error(f"Report generation error: {e}")
        return jsonify({'error': 'Report generation failed'}), 500

def get_write_dir():
    """Return writable directory — /tmp on Vercel, static/ locally."""
    if os.environ.get('VERCEL'):
        os.makedirs('/tmp', exist_ok=True)
        return '/tmp'
    os.makedirs('static', exist_ok=True)
    return 'static'

def download_report_pdf(text, report_data):
    try:
        timestamp = int(time.time())
        filename = f'report_{timestamp}.pdf'
        write_dir = get_write_dir()
        filepath = os.path.join(write_dir, filename)
        
        doc = SimpleDocTemplate(filepath, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor='#2c3e50',
            spaceAfter=30
        )
        story.append(Paragraph(report_data.get('title', 'Summary Report'), title_style))
        story.append(Spacer(1, 0.2*inch))
        
        story.append(Paragraph('<b>Introduction</b>', styles['Heading2']))
        story.append(Paragraph(report_data.get('introduction', ''), styles['BodyText']))
        story.append(Spacer(1, 0.2*inch))
        
        story.append(Paragraph('<b>Key Points</b>', styles['Heading2']))
        for point in report_data.get('key_points', []):
            story.append(Paragraph(f'• {point}', styles['BodyText']))
        story.append(Spacer(1, 0.2*inch))
        
        story.append(Paragraph('<b>Conclusion</b>', styles['Heading2']))
        story.append(Paragraph(report_data.get('conclusion', ''), styles['BodyText']))
        
        doc.build(story)
        
        # On Vercel serve from /tmp via a special route; locally serve from /static
        if os.environ.get('VERCEL'):
            return f'/tmp-file/{filename}'
        return f'/static/{filename}'
    except Exception as e:
        logger.error(f"PDF generation error: {e}")
        return None

@app.route('/')
def index():
    try:
        return render_template('index.html', current_user=current_user)
    except Exception as e:
        logger.error(f"Error rendering index.html: {e}")
        return jsonify({'error': 'Template not found or server error'}), 500

@app.route('/login')
def login_page():
    if current_user.is_authenticated:
        return redirect(url_for('next_page'))
    return render_template('login.html')

@app.route('/login', methods=['POST'])
def login():
    try:
        data = request.get_json()
        email = data.get('email')
        password = data.get('password')
        
        user = User.query.filter_by(email=email).first()
        
        if user and check_password_hash(user.password, password):
            login_user(user, remember=True, duration=timedelta(days=7))
            session.permanent = True
            return jsonify({'success': True, 'message': 'Login successful'})
        else:
            return jsonify({'success': False, 'message': 'Invalid email or password'})
    except Exception as e:
        logger.error(f"Login error: {e}")
        return jsonify({'success': False, 'message': 'Server error'}), 500

@app.route('/signup', methods=['POST'])
def signup():
    try:
        data = request.get_json()
        name = data.get('name')
        email = data.get('email')
        password = data.get('password')
        
        # Check if user already exists
        if User.query.filter_by(email=email).first():
            return jsonify({'success': False, 'message': 'Email already registered'})
        
        # Create new user
        hashed_password = generate_password_hash(password)
        new_user = User(name=name, email=email, password=hashed_password)
        db.session.add(new_user)
        db.session.commit()
        
        return jsonify({'success': True, 'message': 'Account created successfully'})
    except Exception as e:
        logger.error(f"Signup error: {e}")
        return jsonify({'success': False, 'message': 'Server error'}), 500

# Email sending function
def send_otp_email(email, otp):
    try:
        smtp_server = os.getenv('SMTP_SERVER', 'smtp.gmail.com')
        smtp_port = int(os.getenv('SMTP_PORT', 587))
        smtp_email = os.getenv('SMTP_EMAIL')
        smtp_password = os.getenv('SMTP_PASSWORD')
        
        if not smtp_email or not smtp_password:
            logger.warning("SMTP not configured - OTP will be logged to console")
            logger.info(f"=== OTP FOR {email}: {otp} ===")
            print(f"\n{'='*50}")
            print(f"OTP FOR {email}: {otp}")
            print(f"{'='*50}\n")
            return True  # Return True for testing without email
        
        msg = MIMEMultipart('alternative')
        msg['Subject'] = 'Password Reset OTP - Smart Summarizer'
        msg['From'] = smtp_email
        msg['To'] = email
        
        html = f"""
        <html>
          <body style="font-family: Arial, sans-serif; padding: 20px; background-color: #f5f5f5;">
            <div style="max-width: 600px; margin: 0 auto; background: white; padding: 30px; border-radius: 10px; box-shadow: 0 2px 10px rgba(0,0,0,0.1);">
              <h2 style="color: #667eea; margin-bottom: 20px;">Password Reset Request</h2>
              <p style="color: #333; font-size: 16px; line-height: 1.6;">
                You requested to reset your password. Use the OTP below to proceed:
              </p>
              <div style="background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); color: white; font-size: 32px; font-weight: bold; text-align: center; padding: 20px; border-radius: 8px; margin: 30px 0; letter-spacing: 8px;">
                {otp}
              </div>
              <p style="color: #666; font-size: 14px;">
                This OTP will expire in <strong>5 minutes</strong>.
              </p>
              <p style="color: #666; font-size: 14px;">
                If you didn't request this, please ignore this email.
              </p>
              <hr style="border: none; border-top: 1px solid #eee; margin: 30px 0;">
              <p style="color: #999; font-size: 12px; text-align: center;">
                Smart Summarizer - Secure Password Reset
              </p>
            </div>
          </body>
        </html>
        """
        
        part = MIMEText(html, 'html')
        msg.attach(part)
        
        with smtplib.SMTP(smtp_server, smtp_port) as server:
            server.starttls()
            server.login(smtp_email, smtp_password)
            server.send_message(msg)
        
        logger.info(f"OTP email sent to {email}")
        return True
    except Exception as e:
        logger.error(f"Email sending error: {e}")
        # Still log OTP for testing
        logger.info(f"=== OTP FOR {email}: {otp} ===")
        print(f"\n{'='*50}")
        print(f"OTP FOR {email}: {otp}")
        print(f"{'='*50}\n")
        return True  # Return True to allow testing without email

# Rate limiting check
def check_otp_rate_limit(email):
    now = datetime.utcnow()
    one_hour_ago = now - timedelta(hours=1)
    
    # Clean old requests
    otp_request_tracker[email] = [
        ts for ts in otp_request_tracker[email] 
        if ts > one_hour_ago
    ]
    
    # Check if limit exceeded
    if len(otp_request_tracker[email]) >= 3:
        return False
    
    return True

@app.route('/forgot-password/send-otp', methods=['POST'])
def send_otp():
    try:
        data = request.get_json()
        email = data.get('email', '').strip().lower()
        
        if not email:
            return jsonify({'success': False, 'message': 'Email is required'}), 400
        
        # Check if user exists
        user = User.query.filter_by(email=email).first()
        if not user:
            return jsonify({'success': False, 'message': 'Email not found'}), 404
        
        # Check rate limit
        if not check_otp_rate_limit(email):
            return jsonify({'success': False, 'message': 'Too many OTP requests. Please try again after 1 hour.'}), 429
        
        # Generate 6-digit OTP
        otp = ''.join([str(secrets.randbelow(10)) for _ in range(6)])
        
        # Hash OTP before storing
        otp_hash = hashlib.sha256(otp.encode()).hexdigest()
        
        # Store OTP in database
        expires_at = datetime.utcnow() + timedelta(minutes=5)
        new_otp = PasswordResetOTP(
            email=email,
            otp_hash=otp_hash,
            expires_at=expires_at
        )
        db.session.add(new_otp)
        db.session.commit()
        
        # Track request
        otp_request_tracker[email].append(datetime.utcnow())
        
        # Send OTP via email (or log to console if SMTP not configured)
        smtp_configured = os.getenv('SMTP_EMAIL') and os.getenv('SMTP_PASSWORD')
        
        if send_otp_email(email, otp):
            if smtp_configured:
                message = 'OTP sent to your email'
            else:
                message = f'SMTP not configured. Check console for OTP: {otp}'
            
            return jsonify({
                'success': True, 
                'message': message,
                'otp_for_testing': otp if not smtp_configured else None,
                'expires_in': 300  # 5 minutes in seconds
            })
        else:
            return jsonify({'success': False, 'message': 'Failed to send OTP'}), 500
            
    except Exception as e:
        logger.error(f"Send OTP error: {e}")
        return jsonify({'success': False, 'message': 'Server error'}), 500

@app.route('/forgot-password/verify-otp', methods=['POST'])
def verify_otp():
    try:
        data = request.get_json()
        email = data.get('email', '').strip().lower()
        otp = data.get('otp', '').strip()
        
        if not email or not otp:
            return jsonify({'success': False, 'message': 'Email and OTP are required'}), 400
        
        # Hash the provided OTP
        otp_hash = hashlib.sha256(otp.encode()).hexdigest()
        
        # Find valid OTP
        otp_record = PasswordResetOTP.query.filter_by(
            email=email,
            otp_hash=otp_hash,
            is_used=False
        ).first()
        
        if not otp_record:
            return jsonify({'success': False, 'message': 'Invalid OTP'}), 400
        
        # Check if expired
        if datetime.utcnow() > otp_record.expires_at:
            return jsonify({'success': False, 'message': 'OTP has expired'}), 400
        
        # Mark OTP as used
        otp_record.is_used = True
        db.session.commit()
        
        # Generate temporary token for password reset
        reset_token = secrets.token_urlsafe(32)
        session[f'reset_token_{email}'] = {
            'token': reset_token,
            'expires': (datetime.utcnow() + timedelta(minutes=10)).isoformat()
        }
        
        return jsonify({
            'success': True, 
            'message': 'OTP verified successfully',
            'reset_token': reset_token
        })
        
    except Exception as e:
        logger.error(f"Verify OTP error: {e}")
        return jsonify({'success': False, 'message': 'Server error'}), 500

@app.route('/forgot-password/reset-password', methods=['POST'])
def reset_password():
    try:
        data = request.get_json()
        email = data.get('email', '').strip().lower()
        reset_token = data.get('reset_token', '')
        new_password = data.get('new_password', '')
        
        if not email or not reset_token or not new_password:
            return jsonify({'success': False, 'message': 'All fields are required'}), 400
        
        # Validate password strength
        if len(new_password) < 6:
            return jsonify({'success': False, 'message': 'Password must be at least 6 characters'}), 400
        
        # Verify reset token
        session_data = session.get(f'reset_token_{email}')
        if not session_data or session_data['token'] != reset_token:
            return jsonify({'success': False, 'message': 'Invalid reset token'}), 400
        
        # Check token expiry
        if datetime.utcnow() > datetime.fromisoformat(session_data['expires']):
            return jsonify({'success': False, 'message': 'Reset token has expired'}), 400
        
        # Update password
        user = User.query.filter_by(email=email).first()
        if not user:
            return jsonify({'success': False, 'message': 'User not found'}), 404
        
        user.password = generate_password_hash(new_password)
        db.session.commit()
        
        # Clear session
        session.pop(f'reset_token_{email}', None)
        
        return jsonify({
            'success': True, 
            'message': 'Password reset successfully'
        })
        
    except Exception as e:
        logger.error(f"Reset password error: {e}")
        return jsonify({'success': False, 'message': 'Server error'}), 500

@app.route('/logout')
@login_required
def logout():
    logout_user()
    return redirect(url_for('index'))

@app.route('/next')
@login_required
def next_page():
    try:
        return render_template('next.html', current_user=current_user)
    except Exception as e:
        logger.error(f"Error rendering next.html: {e}")
        return jsonify({'error': 'Template not found or server error'}), 500

@app.route('/upload', methods=['POST'])
@login_required
def upload_file():
    try:
        if 'file' not in request.files:
            logger.error("No file part in upload request")
            return jsonify({'error': 'No file part in request'}), 400
        file = request.files['file']
        language = request.form.get('language', 'en')
        if language not in SUPPORTED_LANGUAGES:
            logger.error(f"Unsupported language: {language}")
            return jsonify({'error': f"Unsupported language: {language}"}), 400
        if file.filename == '':
            logger.error("No selected file")
            return jsonify({'error': 'No file selected'}), 400
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        
        # Save file with immediate feedback
        try:
            file.save(file_path)
        except Exception as e:
            logger.error(f"Error saving file {filename}: {e}")
            return jsonify({'error': 'Failed to save file'}), 500
        
        file_type = filename.split('.')[-1].lower()
        
        # Extract text (optimized - limit text size for faster processing)
        text = extract_text_from_file(file_path, file_type, language)
        
        # Clean up file immediately
        try:
            os.remove(file_path)
        except Exception as e:
            logger.warning(f"Error removing file {file_path}: {e}")
        
        if not text:
            logger.error(f"No text extracted from {filename}")
            return jsonify({'error': f"Failed to extract text from {filename}. File may be empty, corrupted, or not supported."}), 400
        
        # Limit text for faster processing (first 5000 chars for summary)
        text_for_summary = text[:5000] if len(text) > 5000 else text
        
        # Generate summary with optimized text
        summary = generate_gemini_summary(text_for_summary, language)
        if not summary:
            logger.error(f"No summary generated for {filename}")
            return jsonify({'error': 'Failed to generate summary'}), 400
        
        return jsonify({'success': True, 'text': text, 'summary': summary})
    except Exception as e:
        logger.error(f"Upload endpoint error: {e}")
        return jsonify({'error': f"Server error during file upload: {str(e)}"}), 500

@app.route('/summarize_url', methods=['POST'])
@login_required
def summarize_url():
    try:
        data = request.get_json()
        url = data.get('url', '')
        language = data.get('language', 'en')
        if language not in SUPPORTED_LANGUAGES:
            logger.error(f"Unsupported language: {language}")
            return jsonify({'error': f"Unsupported language: {language}"}), 400
        if not url:
            logger.error("No URL provided")
            return jsonify({'error': 'No URL provided'}), 400
        if 'youtube.com' in url or 'youtu.be' in url:
            text = extract_youtube_transcript(url, language)
        else:
            try:
                response = requests.get(url, timeout=10)
                response.raise_for_status()
                soup = BeautifulSoup(response.text, 'html.parser')
                text = ' '.join([p.get_text() for p in soup.find_all('p') if p.get_text().strip()])
            except Exception as e:
                logger.error(f"Error fetching URL {url}: {e}")
                text = ''
        if not text:
            logger.error(f"No text extracted from URL {url}")
            return jsonify({'error': 'No text extracted from URL'}), 400
        summary = generate_gemini_summary(text, language)
        if not summary:
            logger.error(f"No summary generated for URL {url}")
            return jsonify({'error': 'Failed to generate summary for URL'}), 400
        return jsonify({'success': True, 'text': text, 'summary': summary})
    except Exception as e:
        logger.error(f"URL summarization error: {e}")
        return jsonify({'error': f"Server error during URL processing: {str(e)}"}), 500

@app.route('/generate_summary', methods=['POST'])
@login_required
def generate_summary():
    try:
        data = request.get_json()
        text = data.get('text', '')
        language = data.get('language', 'en')
        if not text:
            logger.error("No text provided for summary")
            return jsonify({'error': 'No text provided'}), 400
        summary = generate_gemini_summary(text, language)
        if not summary:
            logger.error("No summary generated")
            return jsonify({'error': 'Failed to generate summary'}), 400
        return jsonify({'success': True, 'summary': summary})
    except Exception as e:
        logger.error(f"Summary generation error: {e}")
        return jsonify({'error': f"Server error during summary generation: {str(e)}"}), 500

@app.route('/ask_question', methods=['POST'])
@login_required
def ask_question():
    try:
        data = request.get_json()
        text = data.get('text', '')
        summary = data.get('summary', '')  # Get summary for faster responses
        question = data.get('question', '')
        language = data.get('language', 'en')
        if not text or not question:
            logger.error("Text or question missing")
            return jsonify({'error': 'Text or question missing'}), 400
        answer = generate_gemini_answer(text, question, language, summary)
        if not answer:
            logger.error("No answer generated")
            return jsonify({'error': 'Failed to generate answer'}), 400
        return jsonify({'success': True, 'answer': answer})
    except Exception as e:
        logger.error(f"Question answering error: {e}")
        return jsonify({'error': f"Server error during question answering: {str(e)}"}), 500

@app.route('/generate_mindmap', methods=['POST'])
@login_required
def generate_mindmap():
    try:
        data = request.get_json()
        text = data.get('text', '')
        if not text:
            logger.error("No text provided for mind map")
            return jsonify({'error': 'No text provided'}), 400
        img_base64 = generate_mind_map(text)
        if img_base64:
            return jsonify({
                'success': True, 
                'mindmap_base64': img_base64,
                'filename': f'mindmap_{int(time.time())}.png'
            })
        else:
            return jsonify({'error': 'Failed to generate mind map'}), 400
    except Exception as e:
        logger.error(f"Mind map endpoint error: {e}")
        return jsonify({'error': f"Server error during mind map generation: {str(e)}"}), 500

@app.route('/studio/audio', methods=['POST'])
@login_required
def studio_audio():
    try:
        data = request.get_json()
        text = data.get('text', '')
        language = data.get('language', 'en')
        if not text:
            return jsonify({'error': 'No text provided'}), 400
        return generate_audio_overview(text, language)
    except Exception as e:
        logger.error(f"Studio audio error: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/studio/video', methods=['POST'])
@login_required
def studio_video():
    try:
        data = request.get_json()
        text = data.get('text', '')
        language = data.get('language', 'en')
        if not text:
            return jsonify({'error': 'No text provided'}), 400
        return generate_video_overview(text, language)
    except Exception as e:
        logger.error(f"Studio video error: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/studio/mindmap', methods=['POST'])
@login_required
def studio_mindmap():
    try:
        data = request.get_json()
        text = data.get('text', '')
        if not text:
            return jsonify({'error': 'No text provided'}), 400
        img_base64 = generate_mind_map(text)
        if img_base64:
            return jsonify({
                'success': True, 
                'mindmap_base64': img_base64,
                'filename': f'mindmap_{int(time.time())}.png'
            })
        else:
            return jsonify({'error': 'Failed to generate mind map'}), 400
    except Exception as e:
        logger.error(f"Studio mindmap error: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/studio/report', methods=['POST'])
@login_required
def studio_report():
    try:
        data = request.get_json()
        text = data.get('text', '')
        language = data.get('language', 'en')
        if not text:
            return jsonify({'error': 'No text provided'}), 400
        return generate_report_document(text, language)
    except Exception as e:
        logger.error(f"Studio report error: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/download_report_pdf', methods=['POST'])
@login_required
def download_report_pdf_endpoint():
    try:
        data = request.get_json()
        text = data.get('text', '')
        report_data = data.get('report_data', {})
        
        pdf_path = download_report_pdf(text, report_data)
        if pdf_path:
            return jsonify({'success': True, 'pdf_url': pdf_path})
        else:
            return jsonify({'error': 'Failed to generate PDF'}), 400
    except Exception as e:
        logger.error(f"PDF download error: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/download_summary', methods=['POST'])
@login_required
def download_summary_endpoint():
    try:
        data = request.get_json()
        summary = data.get('summary', '')
        format_type = data.get('format', 'txt')
        
        if not summary:
            return jsonify({'error': 'No summary provided'}), 400
        
        timestamp = int(time.time())
        write_dir = get_write_dir()
        is_vercel = os.environ.get('VERCEL')
        
        if format_type == 'txt':
            filename = f'summary_{timestamp}.txt'
            filepath = os.path.join(write_dir, filename)
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(summary)
            url = f'/tmp-file/{filename}' if is_vercel else f'/static/{filename}'
            return jsonify({'success': True, 'download_url': url})
        
        elif format_type == 'pdf':
            filename = f'summary_{timestamp}.pdf'
            filepath = os.path.join(write_dir, filename)
            doc = SimpleDocTemplate(filepath, pagesize=letter)
            styles = getSampleStyleSheet()
            story = []
            title_style = ParagraphStyle('CustomTitle', parent=styles['Heading1'], fontSize=20, spaceAfter=30)
            story.append(Paragraph('Summary', title_style))
            story.append(Paragraph(summary, styles['BodyText']))
            doc.build(story)
            url = f'/tmp-file/{filename}' if is_vercel else f'/static/{filename}'
            return jsonify({'success': True, 'download_url': url})
        
        elif format_type == 'docx':
            filename = f'summary_{timestamp}.docx'
            filepath = os.path.join(write_dir, filename)
            doc = docx.Document()
            doc.add_heading('Summary', 0)
            doc.add_paragraph(summary)
            doc.save(filepath)
            url = f'/tmp-file/{filename}' if is_vercel else f'/static/{filename}'
            return jsonify({'success': True, 'download_url': url})
        
        else:
            return jsonify({'error': 'Invalid format'}), 400
            
    except Exception as e:
        logger.error(f"Download summary error: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/ping')
def ping():
    groq_key = os.getenv('GROQ_API_KEY')
    return jsonify({
        'success': True, 
        'message': 'Server is running',
        'groq_configured': bool(groq_key),
        'groq_key_prefix': groq_key[:10] if groq_key else None
    })

@app.route('/tmp-file/<filename>')
@login_required
def serve_tmp_file(filename):
    """Serve generated files from /tmp on Vercel."""
    filepath = os.path.join('/tmp', filename)
    if os.path.exists(filepath):
        return send_file(filepath, as_attachment=True)
    return jsonify({'error': 'File not found'}), 404

@app.route('/cleanup')
def cleanup():
    cleanup_old_files()
    return jsonify({'success': True, 'message': 'Cleanup completed'})

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)